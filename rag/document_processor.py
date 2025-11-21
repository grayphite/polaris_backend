"""
Document Processor - Processing for RAG
Extracts text from PDF, Word and other formats
"""

import logging
from typing import Dict, Any, List, Optional
from pathlib import Path

# Safe imports with fallback
try:
	import pypdf
	import pdfplumber
	PDF_AVAILABLE = True
except ImportError:
	PDF_AVAILABLE = False

try:
	from docx import Document as DocxDocument
	DOCX_AVAILABLE = True
except ImportError:
	DOCX_AVAILABLE = False

try:
	import fitz  # PyMuPDF
	PYMUPDF_AVAILABLE = True
except ImportError:
	PYMUPDF_AVAILABLE = False

from .utils import RAGUtils

logger = logging.getLogger(__name__)


class DocumentProcessor:
	"""Legal document processor for RAG"""
	
	def __init__(self):
		self.supported_formats = self._check_supported_formats()
		logger.info(f"DocumentProcessor initialized. Supported formats: {self.supported_formats}")
	
	def _check_supported_formats(self) -> List[str]:
		"""Check which formats are available"""
		formats = ['.txt']  # Always available
		
		if PDF_AVAILABLE:
			formats.extend(['.pdf'])
		
		if DOCX_AVAILABLE:
			formats.extend(['.docx', '.doc'])
		
		return formats
	
	def process_document(self, file_path: str) -> Dict[str, Any]:
		"""
		Process a document and extract text + metadata
		
		Args:
		    file_path: File path
		
		Returns:
		    Dict with text, metadata and chunks
		"""
		try:
			# Validate file
			is_valid, message = RAGUtils.validate_file_path(file_path)
			if not is_valid:
				return {
					'success': False,
					'error': message,
					'text': '',
					'metadata': {},
					'chunks': []
				}
			
			path = Path(file_path)
			extension = path.suffix.lower()
			
			# Extract text based on format
			if extension == '.pdf':
				text = self._extract_pdf_text(file_path)
			elif extension in ['.docx', '.doc']:
				text = self._extract_docx_text(file_path)
			elif extension == '.txt':
				text = self._extract_txt_text(file_path)
			else:
				return {
					'success': False,
					'error': f"Unsupported format: {extension}",
					'text': '',
					'metadata': {},
					'chunks': []
				}
			
			if not text.strip():
				return {
					'success': False,
					'error': "No text extracted from document",
					'text': '',
					'metadata': {},
					'chunks': []
				}
			
			# Clean and normalize text
			text = RAGUtils.clean_text(text)
			
			# Extract metadata
			metadata = RAGUtils.extract_metadata_from_path(file_path)
			metadata.update({
				'text_length': len(text),
				'word_count': len(text.split()),
				'processed_by': 'DocumentProcessor',
				'extraction_method': self._get_extraction_method(extension)
			})
			
			# Generate intelligent chunks
			chunks = RAGUtils.intelligent_chunk_text(text)
			
			# Add metadata to chunks
			for chunk in chunks:
				chunk['source_file'] = path.name
				chunk['source_path'] = file_path
				chunk['document_metadata'] = metadata
			
			logger.info(f"Document processed: {path.name} - {len(chunks)} chunks generated")
			
			return {
				'success': True,
				'text': text,
				'metadata': metadata,
				'chunks': chunks,
				'source_file': path.name
			}
			
		except Exception as e:
			error_msg = f"Error processing document {file_path}: {str(e)}"
			logger.error(error_msg)
			return {
				'success': False,
				'error': error_msg,
				'text': '',
				'metadata': {},
				'chunks': []
			}
	
	def extract_text(self, file_path: str) -> Dict[str, Any]:
		"""
		Alias to process_document. Extracts text from a file.
		Keeps compatibility with tests.
		"""
		return self.process_document(file_path)
	
	def _extract_pdf_text(self, file_path: str) -> str:
		"""Extract text from PDF using multiple libraries"""
		text = ""
		
		# Priority 1: pdfplumber (better for tables/layout)
		if PDF_AVAILABLE:
			try:
				import pdfplumber
				with pdfplumber.open(file_path) as pdf:
					for page_num, page in enumerate(pdf.pages):
						page_text = page.extract_text()
						if page_text:
							text += f"\n\n=== PAGE {page_num + 1} ===\n\n"
							text += page_text
				
				if text.strip():
					logger.debug(f"PDF extracted with pdfplumber: {len(text)} chars")
					return text
			except Exception as e:
				logger.warning(f"pdfplumber failed: {str(e)}")
		
		# Priority 2: PyMuPDF (good overall)
		if PYMUPDF_AVAILABLE:
			try:
				import fitz
				doc = fitz.open(file_path)
				for page_num in range(len(doc)):
					page = doc.load_page(page_num)
					page_text = page.get_text()
					if page_text:
						text += f"\n\n=== PAGE {page_num + 1} ===\n\n"
						text += page_text
				doc.close()
				
				if text.strip():
					logger.debug(f"PDF extracted with PyMuPDF: {len(text)} chars")
					return text
			except Exception as e:
				logger.warning(f"PyMuPDF failed: {str(e)}")
		
		# Priority 3: pypdf (fallback)
		if PDF_AVAILABLE:
			try:
				import pypdf
				with open(file_path, 'rb') as file:
					pdf_reader = pypdf.PdfReader(file)
					for page_num, page in enumerate(pdf_reader.pages):
						page_text = page.extract_text()
						if page_text:
							text += f"\n\n=== PAGE {page_num + 1} ===\n\n"
							text += page_text
				
				if text.strip():
					logger.debug(f"PDF extracted with pypdf: {len(text)} chars")
					return text
			except Exception as e:
				logger.warning(f"pypdf failed: {str(e)}")
		
		raise Exception("No PDF library available or extraction failed")
	
	def _extract_docx_text(self, file_path: str) -> str:
		"""Extract text from Word documents"""
		if not DOCX_AVAILABLE:
			raise Exception("python-docx is not installed")
		
		try:
			from docx import Document
			doc = Document(file_path)
			
			text_parts = []
			
			# Paragraphs
			for paragraph in doc.paragraphs:
				if paragraph.text.strip():
					text_parts.append(paragraph.text)
			
			# Tables
			for table in doc.tables:
				for row in table.rows:
					row_texts = []
					for cell in row.cells:
						if cell.text.strip():
							row_texts.append(cell.text.strip())
					if row_texts:
						text_parts.append(" | ".join(row_texts))
			
			text = "\n\n".join(text_parts)
			logger.debug(f"DOCX extracted: {len(text)} chars")
			return text
			
		except Exception as e:
			raise Exception(f"Error extracting DOCX text: {str(e)}")
	
	def _extract_txt_text(self, file_path: str) -> str:
		"""Extract text from TXT files"""
		try:
			# Try different encodings
			encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
			
			for encoding in encodings:
				try:
					with open(file_path, 'r', encoding=encoding) as file:
						text = file.read()
						logger.debug(f"TXT extracted with {encoding}: {len(text)} chars")
						return text
				except UnicodeDecodeError:
					continue
			
			raise Exception("Could not decode text file")
			
		except Exception as e:
			raise Exception(f"Error extracting TXT text: {str(e)}")
	
	def _get_extraction_method(self, extension: str) -> str:
		"""Return extraction method used"""
		if extension == '.pdf':
			if PDF_AVAILABLE:
				return "pdfplumber/PyMuPDF/pypdf"
			else:
				return "not_available"
		elif extension in ['.docx', '.doc']:
			return "python-docx" if DOCX_AVAILABLE else "not_available"
		elif extension == '.txt':
			return "built-in"
		else:
			return "unsupported"
	
	def process_multiple_documents(self, file_paths: List[str]) -> Dict[str, Any]:
		"""
		Batch processing for multiple documents
		"""
		results = {
			'successful': [],
			'failed': [],
			'total_chunks': 0,
			'total_text_length': 0,
			'processing_summary': {}
		}
		
		for file_path in file_paths:
			try:
				result = self.process_document(file_path)
				
				if result['success']:
					results['successful'].append(result)
					results['total_chunks'] += len(result['chunks'])
					results['total_text_length'] += len(result['text'])
				else:
					results['failed'].append({
						'file_path': file_path,
						'error': result['error']
					})
				
			except Exception as e:
				results['failed'].append({
					'file_path': file_path,
					'error': str(e)
				})
		
		results['processing_summary'] = {
			'total_files': len(file_paths),
			'successful_files': len(results['successful']),
			'failed_files': len(results['failed']),
			'success_rate': len(results['successful']) / len(file_paths) if file_paths else 0,
			'average_chunks_per_file': (results['total_chunks'] / len(results['successful']) 
			                          if results['successful'] else 0),
			'total_text_length': results['total_text_length']
		}
		
		logger.info(f"Batch processing completed: {results['processing_summary']}")
		return results
	
	def get_supported_formats(self) -> List[str]:
		"""Return supported formats list"""
		return self.supported_formats.copy()
	
	def check_dependencies(self) -> Dict[str, bool]:
		"""Return dependency status"""
		return {
			'pypdf': PDF_AVAILABLE,
			'pdfplumber': PDF_AVAILABLE,
			'python-docx': DOCX_AVAILABLE,
			'PyMuPDF': PYMUPDF_AVAILABLE,
			'txt_support': True  # Always available
		}
